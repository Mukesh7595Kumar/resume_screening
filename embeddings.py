"""
This module handles the text extraction, embedding creation, and vector store management.

It provides functions to:
- Extract text from PDF and DOCX files.
- Create a FAISS vector store from the extracted text.
- Use Hugging Face embeddings for vectorization.
"""

import os
import logging
from typing import List, Tuple

import docx
import fitz  # PyMuPDF
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.schema import Document

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')


def extract_text_from_file(file_path: str) -> str:
    """
    Extracts text from a given file (PDF or DOCX).

    Args:
        file_path (str): The path to the file.

    Returns:
        str: The extracted text.
    """
    try:
        if file_path.endswith(".pdf"):
            with fitz.open(file_path) as doc:
                return "".join(page.get_text() for page in doc)
        elif file_path.endswith(".docx"):
            doc = docx.Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        else:
            logging.warning(f"Unsupported file format: {file_path}")
            return ""
    except Exception as e:
        logging.error(f"Error extracting text from {file_path}: {e}")
        return ""


def get_documents_from_folder(folder_path: str) -> List[Document]:
    """
    Extracts text from all supported files in a folder and creates Document objects.

    Args:
        folder_path (str): The path to the folder containing resumes.

    Returns:
        List[Document]: A list of Document objects.
    """
    documents = []
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        if os.path.isfile(file_path):
            text = extract_text_from_file(file_path)
            if text:
                documents.append(Document(page_content=text, metadata={"source": filename}))
    return documents


def create_vector_store(documents: List[Document]) -> FAISS:
    """
    Creates a FAISS vector store from a list of documents.

    Args:
        documents (List[Document]): The list of documents to process.

    Returns:
        FAISS: The created FAISS vector store.
    """
    try:
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        split_documents = text_splitter.split_documents(documents)

        # Use a pre-trained model from Hugging Face for embeddings
        embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

        vector_store = FAISS.from_documents(split_documents, embeddings)
        logging.info("Vector store created successfully.")
        return vector_store
    except Exception as e:
        logging.error(f"Error creating vector store: {e}")
        raise

if __name__ == '__main__':
    # Example usage:
    # Create a dummy resume folder and files for testing
    if not os.path.exists("data/resumes"):
        os.makedirs("data/resumes")
    with open("data/resumes/resume1.txt", "w") as f:
        f.write("This is a sample resume of a Python developer with experience in machine learning.")
    with open("data/resumes/resume2.txt", "w") as f:
        f.write("This is a sample resume of a data scientist with experience in NLP.")

    docs = get_documents_from_folder("data/resumes")
    if docs:
        vector_store = create_vector_store(docs)
        print("Vector store created and ready.")
        # You can now use the vector_store for similarity searches
        results = vector_store.similarity_search("python developer", k=1)
        print(results)
    else:
        print("No documents found to create a vector store.")
